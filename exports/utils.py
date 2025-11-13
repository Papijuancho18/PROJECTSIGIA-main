import os
import json
from django.conf import settings
from django.template.loader import render_to_string
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import openpyxl
from docx import Document

class BaseExporter:
    def __init__(self, report, config):
        self.report = report
        self.config = config
        self.output_dir = os.path.join(settings.MEDIA_ROOT, 'exports')
        os.makedirs(self.output_dir, exist_ok=True)
    
    def export(self):
        raise NotImplementedError

class PDFExporter(BaseExporter):
    def export(self):
        filename = f"{self.report.id}_{self.report.title}.pdf"
        file_path = os.path.join(self.output_dir, filename)
        
        c = canvas.Canvas(file_path, pagesize=letter)
        width, height = letter
        
        # Title
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 50, self.report.title)
        
        # Content
        y_position = height - 100
        c.setFont("Helvetica", 12)
        
        for key, value in self.report.content.items():
            if y_position < 50:
                c.showPage()
                y_position = height - 50
            
            c.drawString(50, y_position, f"{key}: {str(value)[:80]}")
            y_position -= 20
        
        c.save()
        return file_path

def generate_pdf_export(data, config):
    """Generate PDF export from data and configuration"""
    report = data
    config = config
    output_dir = os.path.join(settings.MEDIA_ROOT, 'exports')
    os.makedirs(output_dir, exist_ok=True)
    
    filename = f"{report['id']}_{report['title']}.pdf"
    file_path = os.path.join(output_dir, filename)
    
    c = canvas.Canvas(file_path, pagesize=letter)
    width, height = letter
    
    # Title
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 50, report['title'])
    
    # Content
    y_position = height - 100
    c.setFont("Helvetica", 12)
    
    for key, value in report['content'].items():
        if y_position < 50:
            c.showPage()
            y_position = height - 50
        
        c.drawString(50, y_position, f"{key}: {str(value)[:80]}")
        y_position -= 20
    
    c.save()
    return file_path

class WordExporter(BaseExporter):
    def export(self):
        filename = f"{self.report.id}_{self.report.title}.docx"
        file_path = os.path.join(self.output_dir, filename)
        
        doc = Document()
        doc.add_heading(self.report.title, 0)
        
        if self.report.description:
            doc.add_paragraph(self.report.description)
        
        for key, value in self.report.content.items():
            doc.add_heading(key, level=1)
            doc.add_paragraph(str(value))
        
        doc.save(file_path)
        return file_path

def generate_docx_export(data, config):
    """Generate DOCX export from data and configuration"""
    report = data
    config = config
    output_dir = os.path.join(settings.MEDIA_ROOT, 'exports')
    os.makedirs(output_dir, exist_ok=True)
    
    filename = f"{report['id']}_{report['title']}.docx"
    file_path = os.path.join(output_dir, filename)
    
    doc = Document()
    doc.add_heading(report['title'], 0)
    
    if report.get('description'):
        doc.add_paragraph(report['description'])
    
    for key, value in report['content'].items():
        doc.add_heading(key, level=1)
        doc.add_paragraph(str(value))
    
    doc.save(file_path)
    return file_path

class ExcelExporter(BaseExporter):
    def export(self):
        filename = f"{self.report.id}_{self.report.title}.xlsx"
        file_path = os.path.join(self.output_dir, filename)
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Report Data"
        
        # Headers
        ws['A1'] = "Field"
        ws['B1'] = "Value"
        
        # Data
        row = 2
        for key, value in self.report.content.items():
            ws[f'A{row}'] = key
            ws[f'B{row}'] = str(value)
            row += 1
        
        wb.save(file_path)
        return file_path

def generate_xlsx_export(data, config):
    """Generate XLSX export from data and configuration"""
    report = data
    config = config
    output_dir = os.path.join(settings.MEDIA_ROOT, 'exports')
    os.makedirs(output_dir, exist_ok=True)
    
    filename = f"{report['id']}_{report['title']}.xlsx"
    file_path = os.path.join(output_dir, filename)
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Report Data"
    
    # Headers
    ws['A1'] = "Field"
    ws['B1'] = "Value"
    
    # Data
    row = 2
    for key, value in report['content'].items():
        ws[f'A{row}'] = key
        ws[f'B{row}'] = str(value)
        row += 1
    
    wb.save(file_path)
    return file_path

class HTMLExporter(BaseExporter):
    def export(self):
        filename = f"{self.report.id}_{self.report.title}.html"
        file_path = os.path.join(self.output_dir, filename)
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>{self.report.title}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ color: #333; }}
                .content {{ margin: 20px 0; }}
            </style>
        </head>
        <body>
            <h1>{self.report.title}</h1>
            <p>{self.report.description}</p>
            <div class="content">
        """
        
        for key, value in self.report.content.items():
            html_content += f"<h2>{key}</h2><p>{value}</p>"
        
        html_content += """
            </div>
        </body>
        </html>
        """
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return file_path

def generate_html_export(data, config):
    """Generate HTML export from data and configuration"""
    report = data
    config = config
    output_dir = os.path.join(settings.MEDIA_ROOT, 'exports')
    os.makedirs(output_dir, exist_ok=True)
    
    filename = f"{report['id']}_{report['title']}.html"
    file_path = os.path.join(output_dir, filename)
    
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>{report['title']}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; }}
            h1 {{ color: #333; }}
            .content {{ margin: 20px 0; }}
        </style>
    </head>
    <body>
        <h1>{report['title']}</h1>
        <p>{report.get('description', '')}</p>
        <div class="content">
    """
    
    for key, value in report['content'].items():
        html_content += f"<h2>{key}</h2><p>{value}</p>"
    
    html_content += """
        </div>
    </body>
    </html>
    """
    
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    return file_path

class JSONExporter(BaseExporter):
    def export(self):
        filename = f"{self.report.id}_{self.report.title}.json"
        file_path = os.path.join(self.output_dir, filename)
        
        data = {
            'id': self.report.id,
            'title': self.report.title,
            'description': self.report.description,
            'content': self.report.content,
            'status': self.report.status,
            'created_at': self.report.created_at.isoformat(),
            'updated_at': self.report.updated_at.isoformat(),
        }
        
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        return file_path
